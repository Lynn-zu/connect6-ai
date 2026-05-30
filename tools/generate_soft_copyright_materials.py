from __future__ import annotations

import os
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "soft_copyright_materials"
SOFTWARE_NAME = "六子棋智能博弈系统"
VERSION = "V1.0"
TODAY = datetime.now().strftime("%Y年%m月%d日")

SOURCE_FILES = [
    Path("app.py"),
    Path("train.py"),
    Path("templates/index.html"),
    Path("tests/test_ai_regression.py"),
    Path("requirements.txt"),
    Path("README.md"),
]


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        f.write(content.replace("\n", "\r\n"))


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_title(document: Document, title: str, subtitle: str = "") -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    if subtitle:
        sp = document.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run(subtitle)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_manual_docx(md_sections: list[tuple[str, list[str]]], path: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, f"{SOFTWARE_NAME}软件说明书", f"{VERSION} | 申请材料草稿 | {TODAY}")

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    for heading, paragraphs in md_sections:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            if paragraph.startswith("- "):
                doc.add_paragraph(paragraph[2:], style="List Bullet")
            elif paragraph.startswith("1. ") or paragraph.startswith("2. ") or paragraph.startswith("3. ") or paragraph.startswith("4. ") or paragraph.startswith("5. "):
                doc.add_paragraph(paragraph[3:], style="List Number")
            else:
                doc.add_paragraph(paragraph)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_text_docx(title: str, text: str, path: Path, code: bool = False) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, title, f"{VERSION} | 申请材料草稿 | {TODAY}")
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    if code:
        style = doc.styles["Normal"]
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(7.5)
        lines = text.splitlines()
        body_started = False
        for line in lines:
            if line.startswith("===== 第 ") and body_started:
                doc.add_page_break()
            if line.startswith("===== 第 "):
                body_started = True
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(7.5)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
    else:
        for raw in text.splitlines():
            if raw.startswith("# "):
                doc.add_heading(raw[2:], level=1)
            elif raw.startswith("## "):
                doc.add_heading(raw[3:], level=2)
            elif raw.startswith("### "):
                doc.add_heading(raw[4:], level=3)
            elif raw.startswith("- "):
                doc.add_paragraph(raw[2:], style="List Bullet")
            elif raw and raw[0].isdigit() and ". " in raw[:4]:
                doc.add_paragraph(raw.split(". ", 1)[1], style="List Number")
            elif raw.strip():
                doc.add_paragraph(raw)
            else:
                doc.add_paragraph()

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def make_source_identification() -> str:
    numbered: list[str] = []
    for rel in SOURCE_FILES:
        full = ROOT / rel
        if not full.exists():
            continue
        numbered.append(f"// ===== 文件：{rel.as_posix()} =====")
        for idx, line in enumerate(read_text(full).splitlines(), start=1):
            numbered.append(f"{rel.as_posix()}:{idx:04d}  {line}")

    total = len(numbered)
    if total <= 3000:
        selected = numbered
        rule = f"源程序总行数 {total} 行，未超过 3000 行，提交全部源程序。"
    else:
        selected = numbered[:1500] + numbered[-1500:]
        rule = f"源程序总行数 {total} 行，超过 3000 行，按常见鉴别材料格式提交前 1500 行和后 1500 行。"

    pages: list[str] = []
    for page_index in range(0, len(selected), 50):
        page_no = page_index // 50 + 1
        pages.append(f"\n===== 第 {page_no:02d} 页 / 每页约 50 行 =====")
        pages.extend(selected[page_index:page_index + 50])

    header = [
        f"{SOFTWARE_NAME} {VERSION} 源程序鉴别材料",
        f"生成日期：{TODAY}",
        rule,
        "说明：本文件仅整理项目自有源代码、前端页面、测试与依赖清单，不包含虚拟环境、缓存文件、训练数据和模型权重。",
        "",
    ]
    return "\n".join(header + pages) + "\n"


def make_source_manifest() -> str:
    rows = []
    for rel in SOURCE_FILES:
        full = ROOT / rel
        if not full.exists():
            continue
        text = read_text(full)
        rows.append((rel.as_posix(), len(text.splitlines()), full.stat().st_size))

    lines = [
        f"# {SOFTWARE_NAME}源文件目录清单",
        "",
        f"- 软件版本：{VERSION}",
        f"- 生成日期：{TODAY}",
        "- 说明：以下文件为本次软著材料选用的主要源程序文件。",
        "",
        "| 序号 | 文件路径 | 行数 | 大小（字节） | 用途 |",
        "| --- | --- | ---: | ---: | --- |",
    ]
    purpose = {
        "app.py": "Flask 后端、六子棋规则、AI 对弈接口、模型加载与训练接口",
        "train.py": "强化学习模型结构、自博弈训练、MCTS、战术样本与教师样本训练",
        "templates/index.html": "Web 前端界面、棋盘绘制、交互控制、接口调用",
        "tests/test_ai_regression.py": "AI 行为回归测试",
        "requirements.txt": "运行依赖清单",
        "README.md": "项目启动与训练说明",
    }
    for idx, (name, lines_count, size) in enumerate(rows, start=1):
        lines.append(f"| {idx} | `{name}` | {lines_count} | {size} | {purpose.get(name, '项目文件')} |")
    return "\n".join(lines) + "\n"


def make_manual_sections() -> list[tuple[str, list[str]]]:
    return [
        ("一、软件概述", [
            f"{SOFTWARE_NAME}是一套基于 Python Flask 和 Web 前端实现的六子棋人机博弈系统。系统围绕六子棋规则、棋盘交互、传统搜索 AI、强化学习模型 AI、模型训练与对局复盘等功能展开，面向棋类博弈教学、算法实验和智能对弈演示场景。",
            "系统采用浏览器作为操作界面，后端负责规则判定、局面维护、AI 落子计算和模型管理。用户可选择不同 AI 类型、棋盘规格和执棋方，完成开局、落子、悔棋、前进、重新开局、查看 AI 状态等操作。",
        ]),
        ("二、运行环境", [
            "- 操作系统：Windows、Linux 或 macOS。",
            "- 运行语言：Python 3.10/3.11。",
            "- 后端框架：Flask、Flask-CORS。",
            "- 算法与模型：NumPy、TensorFlow/Keras、h5py。",
            "- 浏览器：Chrome、Edge、Firefox 等现代浏览器。",
            "- 默认模型路径：models/connect6.model.h5。",
        ]),
        ("三、主要功能", [
            "- 新建对局：支持设置棋盘尺寸、执棋方和 AI 类型。",
            "- 六子棋规则控制：黑方首手一子，之后双方每回合两子；系统自动检查合法落子、胜负和棋盘状态。",
            "- 人机对弈：支持传统 Minimax 搜索 AI、强化学习增强 AI、纯模型 AI 三类模式。",
            "- 棋盘交互：通过 Canvas 绘制棋盘和棋子，支持鼠标和触摸操作。",
            "- 悔棋与前进：保存落子历史，可撤销或恢复历史步骤。",
            "- 对局状态展示：显示当前回合、胜负结果、AI 类型、模型加载状态和 AI 决策状态。",
            "- 模型训练接口：支持从页面或接口启动后台训练，查询训练状态，查看模型信息并下载模型。",
            "- 复盘能力：后端保留历史棋步，可按步骤重建棋局状态。",
        ]),
        ("四、使用流程", [
            "1. 安装 Python 运行环境并进入项目目录。",
            "2. 根据 requirements.txt 安装依赖。",
            "3. 执行 python app.py 启动服务。",
            "4. 在浏览器访问 http://localhost:5000。",
            "5. 在开始界面选择 AI 类型、棋盘大小和执棋方后进入游戏。",
            "6. 在棋盘上点击交叉点完成落子，系统将在玩家回合结束后自动触发 AI 落子。",
            "7. 对局过程中可使用悔棋、前进、重新开局和更换机器人功能。",
        ]),
        ("五、接口说明", [
            "- GET /：返回六子棋 Web 操作页面。",
            "- POST /api/new_game：创建新对局。",
            "- POST /api/move：提交玩家落子并触发 AI 回合。",
            "- POST /api/undo：撤销最近回合。",
            "- POST /api/redo：恢复被撤销的回合。",
            "- POST /api/restart：重置当前对局。",
            "- GET /api/game_state：获取当前棋局状态。",
            "- POST /api/review/toggle、/api/review/previous、/api/review/next：控制复盘状态。",
            "- POST /api/train 或 /api/train_model：启动模型训练任务。",
            "- GET /api/train_status：查询训练状态。",
            "- GET /api/model_info：查询模型文件与模型结构信息。",
            "- GET /api/download_model：下载当前模型文件。",
        ]),
        ("六、输入输出", [
            "系统输入包括用户选择的棋盘尺寸、执棋方、AI 模式、棋盘落点坐标和训练参数。系统输出包括当前棋盘矩阵、胜负状态、下一步落子提示、AI 落子结果、训练状态、模型状态和模型文件下载结果。",
        ]),
        ("七、异常处理", [
            "- 当 TensorFlow 不可用或模型文件不存在时，系统显示模型不可用状态，并可切换到传统 AI 或启发式策略。",
            "- 当用户落子数量不符合规则、落点越界或位置已被占用时，后端返回错误信息。",
            "- 当训练任务已在运行时，系统避免重复启动训练。",
            "- 当模型训练产生异常或模型权重出现 NaN/Inf 时，训练脚本拒绝覆盖主模型文件。",
        ]),
    ]


def markdown_from_sections(title: str, sections: list[tuple[str, list[str]]]) -> str:
    lines = [f"# {title}", "", f"- 软件名称：{SOFTWARE_NAME}", f"- 版本号：{VERSION}", f"- 生成日期：{TODAY}", ""]
    for heading, paragraphs in sections:
        lines.extend([f"## {heading}", ""])
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
    return "\n".join(lines)


def make_design_doc() -> str:
    return f"""# {SOFTWARE_NAME}设计说明书

- 软件名称：{SOFTWARE_NAME}
- 版本号：{VERSION}
- 生成日期：{TODAY}

## 一、总体设计

系统采用浏览器前端、Flask 后端和本地模型文件组成的轻量化 Web 架构。前端负责界面呈现、棋盘绘制和用户操作采集；后端负责棋局规则、状态维护、AI 决策、模型加载、训练任务调度和接口输出。训练脚本独立承担强化学习模型训练、样本生成、模型校验和模型保存。

## 二、模块划分

1. 棋盘规则模块：由 Connect6Board 实现棋盘初始化、合法性检查、落子、胜负判断、历史记录、撤销、重放和状态序列化。
2. 传统 AI 模块：由 Connect6MinimaxAI 实现局面评估、候选点筛选、Alpha-Beta 搜索、即时胜利检测和防守判断。
3. 强化学习 AI 模块：由 Connect6DeepLearningAI 和 Connect6PureModelAI 实现模型推理、策略输出、候选点筛选、战术补强和降级处理。
4. Web 服务模块：由 Flask 路由提供新局、落子、悔棋、复盘、训练、模型信息和模型下载接口。
5. 前端交互模块：由 templates/index.html 实现页面布局、模式选择、Canvas 棋盘绘制、事件监听和后端接口调用。
6. 模型训练模块：由 train.py 实现策略价值网络、棋盘特征编码、MCTS、自博弈、教师样本、战术样本、模型诊断和模型持久化。

## 三、关键数据结构

- 棋盘矩阵：使用二维数组表示棋盘状态，0 表示空位，1 表示黑棋，2 表示白棋。
- 落子历史：按回合保存落子坐标和玩家编号，用于悔棋、前进和复盘。
- 模型输入：将棋盘转换为 15x15x17 的特征张量，提供给策略价值网络。
- 模型输出：策略头输出棋盘各点概率，价值头输出当前局面对当前玩家的价值估计。
- 训练元数据：保存累计自博弈局数、最近训练时间、参数配置、诊断结果和训练历史。

## 四、核心流程

### 1. 对局流程

用户在前端选择设置后调用新局接口，后端创建棋盘和 AI 实例。玩家落子时，前端提交坐标，后端校验规则并更新棋盘；若未分出胜负且轮到 AI，后端调用对应 AI 生成落子并再次更新状态，最后将完整棋局状态返回前端刷新页面。

### 2. AI 决策流程

系统优先判断即时取胜和必须防守的战术点；强化学习增强模式在模型可用时读取策略分布并结合位置评分选择落点；纯模型模式直接依据模型策略选取合法点；当模型不可用或棋盘规格不匹配时，系统回退至启发式或传统搜索策略。

### 3. 训练流程

训练脚本加载现有模型或创建新模型，通过自博弈和样本生成构造训练数据，执行策略价值网络训练。训练后进行权重有效性检查、策略分布诊断和战术响应验证，验证通过后保存模型并更新训练元数据。

## 五、技术特点

- 采用 Flask 提供简洁的本地 Web 服务，部署和调试成本低。
- 前端使用 Canvas 绘制棋盘，可适配不同屏幕尺寸和触摸操作。
- 同时支持传统搜索 AI 与神经网络模型 AI，便于算法对比。
- 训练接口与对弈系统集成，可在系统内查看训练和模型状态。
- 模型保存前设置多重校验，降低异常训练覆盖有效模型的风险。

## 六、部署说明

安装依赖后执行 python app.py 启动服务，浏览器访问 http://localhost:5000 即可使用。若需要强化学习模式，应保证 models/connect6.model.h5 存在且 TensorFlow 可正常加载；若缺少模型，系统仍可通过传统 AI 模式运行。
"""


def make_application_info() -> str:
    return f"""# 软件著作权登记信息填写参考

> 说明：以下内容是根据当前项目整理的填写草稿，带有【请填写】的项目需要由申请人确认后再提交。

## 基本信息

- 软件全称：{SOFTWARE_NAME}
- 软件简称：六子棋博弈系统
- 版本号：{VERSION}
- 软件分类：游戏软件 / 人工智能应用软件 / 教学实验软件（按申请系统可选项择一）
- 开发完成日期：【请填写，建议填实际完成日期，例如 2026-05-29】
- 首次发表日期：【未发表可填“未发表”；已公开则填写首次公开日期】
- 开发方式：独立开发
- 权利取得方式：原始取得
- 权利范围：全部权利
- 著作权人：【请填写姓名或单位名称】
- 身份证号或统一社会信用代码：【请填写】
- 联系人、电话、邮箱、地址：【请填写】

## 软件功能摘要

{SOFTWARE_NAME}是一套基于 Python Flask 和 Web 前端实现的六子棋人机博弈系统，支持六子棋规则判定、浏览器棋盘交互、传统 Minimax 搜索 AI、强化学习模型 AI、对局状态管理、悔棋前进、复盘、模型训练、模型状态查询和模型下载等功能。

## 技术特点摘要

系统采用前后端分离式交互思路，前端使用 HTML/CSS/JavaScript 与 Canvas 完成棋盘绘制和用户操作，后端使用 Flask 提供接口服务，算法层使用 NumPy、TensorFlow/Keras 构建策略价值网络，并结合 MCTS、自博弈训练、教师样本和战术样本完成模型训练与校验。

## 硬件环境

普通 PC 或服务器，建议内存 8GB 以上；如需训练强化学习模型，建议使用具备较好 CPU/GPU 性能的设备。

## 软件环境

Python 3.10/3.11，Flask，Flask-CORS，NumPy，TensorFlow/Keras，h5py，tqdm，现代浏览器。
"""


def make_checklist() -> str:
    return f"""# {SOFTWARE_NAME}软著提交材料清单

## 已准备文件

1. `01_软件著作权登记信息填写参考.md`：登记表填写草稿。
2. `02_六子棋智能博弈系统_软件说明书.md`：软件说明书 Markdown 版。
3. `02_六子棋智能博弈系统_软件说明书.docx`：软件说明书 Word 版。
4. `03_六子棋智能博弈系统_设计说明书.md`：软件设计说明书。
5. `03_六子棋智能博弈系统_设计说明书.docx`：软件设计说明书 Word 版。
6. `04_六子棋智能博弈系统_源程序鉴别材料.txt`：源程序鉴别材料，按前后各 1500 行整理。
7. `04_六子棋智能博弈系统_源程序鉴别材料.docx`：源程序鉴别材料 Word 版。
8. `05_源文件目录清单.md`：本次选用源文件目录与用途说明。
9. `06_六子棋智能博弈系统_源码归档.zip`：项目源码归档，不含虚拟环境、缓存、训练数据和模型权重备份。

## 仍需申请人补充

1. 著作权人姓名或单位名称。
2. 身份证号或统一社会信用代码。
3. 联系人、手机号、邮箱、通信地址。
4. 开发完成日期、首次发表日期或未发表声明。
5. 如委托代理办理，还需代理委托书和代理人信息。
6. 如多人合作开发，还需合作开发协议或权属说明。

## 提交建议

- 源程序鉴别材料通常要求页眉或首页标注软件名称和版本号；本材料已在文件开头标注。
- 说明书应和源代码功能保持一致，避免写入当前系统没有实现的功能。
- 正式提交前建议将所有【请填写】占位符替换为真实信息。
- 不建议提交 `.venv`、`__pycache__`、`.idea`、训练数据缓存和历史模型备份。
"""


def make_zip(path: Path) -> None:
    include = [
        Path("app.py"),
        Path("train.py"),
        Path("README.md"),
        Path("requirements.txt"),
        Path("templates/index.html"),
        Path("tests/test_ai_regression.py"),
    ]
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in include:
            full = ROOT / rel
            if full.exists():
                zf.write(full, arcname=rel.as_posix())


def main() -> None:
    OUT.mkdir(exist_ok=True)

    manual_sections = make_manual_sections()
    write_text(OUT / "00_提交材料清单.md", make_checklist())
    write_text(OUT / "01_软件著作权登记信息填写参考.md", make_application_info())
    write_text(OUT / "02_六子棋智能博弈系统_软件说明书.md", markdown_from_sections(f"{SOFTWARE_NAME}软件说明书", manual_sections))
    build_manual_docx(manual_sections, OUT / "02_六子棋智能博弈系统_软件说明书.docx")
    design_doc = make_design_doc()
    write_text(OUT / "03_六子棋智能博弈系统_设计说明书.md", design_doc)
    build_text_docx(f"{SOFTWARE_NAME}设计说明书", design_doc, OUT / "03_六子棋智能博弈系统_设计说明书.docx")
    source_doc = make_source_identification()
    write_text(OUT / "04_六子棋智能博弈系统_源程序鉴别材料.txt", source_doc)
    build_text_docx(f"{SOFTWARE_NAME}源程序鉴别材料", source_doc, OUT / "04_六子棋智能博弈系统_源程序鉴别材料.docx", code=True)
    write_text(OUT / "05_源文件目录清单.md", make_source_manifest())
    make_zip(OUT / "06_六子棋智能博弈系统_源码归档.zip")


if __name__ == "__main__":
    main()
